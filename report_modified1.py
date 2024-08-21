import os
import csv
from datetime import datetime
from docx2pdf import convert
from PIL import Image, ExifTags
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


report_created = datetime.now()
message = "st" if "%d" == 1 or "%d" == 21 or "%d" == 31 else "nd" if "%d" == 2 or "%d" == 22 else "rd" if "%d" == 3 or "%d" == 23 else "th"
report_created_time = report_created.strftime("%B %d{}, %Y; %H:%M:%S".format(message))

main_folder = os.getcwd()
main_folder_contents = os.listdir(main_folder)
logo = "logo.png"

folder_count = 1
document = Document()

for content in main_folder_contents:
    if os.path.isdir(content):
        sub_folder = os.listdir(content)
        for file in sub_folder:
            if file.endswith("JPG"):
                path_1 = os.path.join(content, file)
                parent_image = path_1
                parent_file_name = os.path.basename(path_1)
                parent_image_file = Image.open(path_1)
                exif_parent = {ExifTags.TAGS[tag]: value for tag,
                               value in parent_image_file._getexif().items() if tag in ExifTags.TAGS}
                if "GPSInfo" in exif_parent:
                    latitude = "{0}° {1}\' {2}\" {3}".format(
                        round(exif_parent["GPSInfo"][2][0], 0),
                        round(exif_parent["GPSInfo"][2][1], 0),
                        exif_parent["GPSInfo"][2][2],
                        exif_parent["GPSInfo"][1]
                    )
                    longitude = "{0}° {1}\' {2}\" {3}".format(
                        round(exif_parent["GPSInfo"][4][0], 0),
                        round(exif_parent["GPSInfo"][4][1], 0),
                        exif_parent["GPSInfo"][4][2],
                        exif_parent["GPSInfo"][3]
                    )
                    altitude = exif_parent["GPSInfo"][6]
            elif file.endswith("csv"):
                fault_type, severity = list(), list()

                path = os.path.join(content, file)
                with open(path, newline='') as csvfile:
                    fault_reader = csv.reader(csvfile, delimiter=",")
                    for row in fault_reader:
                        if "class" in row or "" in row:
                            continue
                        fault_type.append(row[1])
                        severity.append(row[6])

            elif file.endswith("png"):
                path_2 = os.path.join(content, file)
                child_image_file = path_2
                child_file_name = os.path.basename(path_2)

                if folder_count == 1:
                    document.add_picture(logo, width=Inches(1.25))
                    document.add_heading("INFOGRAPH", 0)
                    document.add_heading("POWERLINE INSPECTION SYSTEM REPORT", 0)
                    document.add_heading("Report Created: " +
                                        report_created_time, 4)

                    first_page = (document.paragraphs[0],
                                document.paragraphs[1], document.paragraphs[2],
                                document.paragraphs[3]
                                )

                    for paragraph in first_page:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    document.add_page_break()

                    p = document.add_paragraph()
                    p.add_run("Original Image Name: ").bold = True
                    p.add_run(parent_file_name)
                    p = document.add_paragraph()
                    p.add_run("Image Path: ").bold = True
                    p.add_run(os.path.join(main_folder, path_1))
                    


                    records = (
                        ("Date Captured", exif_parent["DateTime"]),
                        ("Latitude", latitude),
                        ("Longitude", longitude),
                        ("Altitude", str(altitude) + " m")

                    )

                    table = document.add_table(rows=1, cols=2)
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = 'Image Info'
                    hdr_cells[1].text = 'Details'

                    for i_info, detail in records:
                        row_cells = table.add_row().cells
                        row_cells[0].text = i_info
                        row_cells[1].text = detail
                    table.style = 'Light Shading'

                    p = document.add_paragraph()

                    document.add_heading("FAULT", 1)
                    p = document.add_paragraph()
                    p.add_run( "Fault Image Name: ").bold = True
                    p.add_run(child_file_name)
                    p = document.add_paragraph()
                    p.add_run("Image Path: ").bold = True
                    p.add_run(os.path.join(main_folder, path_2))

                    records1 = (
                        (fault_type[0], severity[0]+" %"),
                        (fault_type[1], severity[1]+" %"),
                        (fault_type[2], severity[2]+" %")
                    )

                    table1 = document.add_table(rows=1, cols=2)
                    hdr_cells = table1.rows[0].cells
                    hdr_cells[0].text = 'Fault Type'
                    hdr_cells[1].text = 'Severity'

                    for f_type, status in records1:
                        row_cells = table1.add_row().cells
                        row_cells[0].text = f_type
                        row_cells[1].text = status
                    table1.style = 'Light Shading'

                    q = document.add_paragraph()

                    document.add_picture(child_image_file, width=Inches(4))
                    child_image = document.paragraphs[-1]
                    child_image.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    folder_count += 1
                    
                else:
                    document.add_page_break()
                    p = document.add_paragraph()
                    p.add_run("Original Image Name: ").bold = True
                    p.add_run(parent_file_name)
                    p = document.add_paragraph()
                    p.add_run("Image Path: ").bold = True
                    p.add_run(os.path.join(main_folder, path_1))
                    

                    records = (
                        ("Date Captured", exif_parent["DateTime"]),
                        ("Latitude", latitude),
                        ("Longitude", longitude),
                        ("Altitude", str(altitude) + " m")

                    )

                    table = document.add_table(rows=1, cols=2)
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = 'Image Info'
                    hdr_cells[1].text = 'Details'

                    for i_info, detail in records:
                        row_cells = table.add_row().cells
                        row_cells[0].text = i_info
                        row_cells[1].text = detail
                    table.style = 'Light Shading'

                    p = document.add_paragraph()

                    document.add_heading("FAULT", 1)
                    p = document.add_paragraph()
                    p.add_run("Fault Image Name: ").bold = True
                    p.add_run(child_file_name)
                    p = document.add_paragraph()
                    p.add_run("Image Path: ").bold = True
                    p.add_run(os.path.join(main_folder, path_2))

                    records1 = (
                        (fault_type[0], severity[0]+" %"),
                        (fault_type[1], severity[1]+" %"),
                        (fault_type[2], severity[2]+" %")
                    )

                    table1 = document.add_table(rows=1, cols=2)
                    hdr_cells = table1.rows[0].cells
                    hdr_cells[0].text = 'Fault Type'
                    hdr_cells[1].text = 'Severity'

                    for f_type, status in records1:
                        row_cells = table1.add_row().cells
                        row_cells[0].text = f_type
                        row_cells[1].text = status
                    table1.style = 'Light Shading'

                    q = document.add_paragraph()

                    document.add_picture(child_image_file, width=Inches(4))
                    child_image = document.paragraphs[-1]
                    child_image.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    section = document.sections[0]
                    footer = section.footer
                    footer_paragraph = footer.paragraphs[0]
                    footer_paragraph.text = "\t\tpowerline inspection system report"
                    
                document_name = "Report.docx"
                document.save(document_name)
convert(document_name)
